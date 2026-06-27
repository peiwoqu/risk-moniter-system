"""
Agent智能风险监控管理系统 - 报告生成服务
生成Word/Excel/PDF格式的风险分析报告
"""

import json
import os
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class ReportService:
    """报告生成服务"""

    BASE_DIR = os.path.dirname(os.path.dirname(__file__))
    REPORT_DIR = os.path.join(BASE_DIR, 'data', 'reports')

    def __init__(self):
        os.makedirs(self.REPORT_DIR, exist_ok=True)

    def generate_summary(self, risks: List[Dict],
                         stats: Dict[str, Any]) -> str:
        """生成风险摘要（300字）"""
        total = stats.get('total', len(risks))
        by_level = stats.get('by_level', {})
        critical = by_level.get('critical', 0)
        high = by_level.get('high', 0)
        avg_score = stats.get('average_score', 0)

        # 找出得分最高的风险
        sorted_risks = sorted(risks,
                            key=lambda r: r.get('risk_score', 0),
                            reverse=True)
        top_risks = sorted_risks[:3]

        top_risk_names = "、".join([r.get('name', '') for r in top_risks])

        summary = (
            f"本报告基于Agent智能风险监控管理系统自动生成，对腾讯控股（0700.HK）"
            f"进行全面的企业风险管理分析。系统共识别{total}项主要风险，其中"
            f"重大风险{critical}项、高度关注风险{high}项，平均风险得分{avg_score}分。"
            f"得分最高的风险为：{top_risk_names}。"
            f"系统基于COSO ERM 2017框架，通过多Agent协作完成了风险识别、"
            f"量化评估（P×I矩阵）、预警检测和应对建议生成等完整流程。"
            f"针对识别出的重大风险，系统提出了基于4T策略（规避/降低/转移/接受）"
            f"的应对方案和改进建议，为企业风险管理决策提供数据驱动支撑。"
        )

        return summary

    def generate_word_report(self,
                             enterprise_name: str,
                             risks: List[Dict],
                             stats: Dict[str, Any],
                             matrix_data: Dict[str, Any],
                             warnings: List[Dict],
                             strategy_summary: Dict[str, Any],
                             priority_actions: List[Dict]) -> str:
        """
        生成Word格式的完整报告

        Returns:
            报告文件路径
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx未安装")
            return ""

        doc = Document()

        # 标题
        title = doc.add_heading(f'{enterprise_name}智能风险监控报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元数据
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}\n')
        meta.add_run(f'分析引擎：Agent智能风险监控管理系统 v1.0\n')
        meta.add_run(f'理论框架：COSO ERM 2017 / ISO 31000')

        doc.add_page_break()

        # 1. 摘要
        doc.add_heading('一、摘要', level=1)
        summary_text = self.generate_summary(risks, stats)
        doc.add_paragraph(summary_text)

        # 2. 风险概览
        doc.add_heading('二、风险概览', level=1)

        # 统计表
        doc.add_heading('2.1 风险统计', level=2)
        table = doc.add_table(rows=1, cols=2, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text = '指标'
        hdr[1].text = '数值'

        stat_rows = [
            ('风险总数', str(stats.get('total', 0))),
            ('重大风险（红色）', str(stats.get('by_level', {}).get('critical', 0))),
            ('高度关注风险（橙色）', str(stats.get('by_level', {}).get('high', 0))),
            ('一般风险（黄色）', str(stats.get('by_level', {}).get('medium', 0))),
            ('低风险（绿色）', str(stats.get('by_level', {}).get('low', 0))),
            ('平均风险得分', str(stats.get('average_score', 0))),
            ('上升趋势风险', str(stats.get('by_trend', {}).get('increasing', 0))),
        ]
        for label, value in stat_rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        # 风险清单表
        doc.add_heading('2.2 风险清单', level=2)
        risk_table = doc.add_table(rows=1, cols=7, style='Table Grid')
        risk_hdr = risk_table.rows[0].cells
        headers = ['编号', '风险名称', '类别', 'P', 'I', '得分', '等级']
        for i, h in enumerate(headers):
            risk_hdr[i].text = h

        level_map = {"critical": "重大", "high": "高度关注",
                    "medium": "一般", "low": "低"}

        for risk in risks:
            row = risk_table.add_row()
            row.cells[0].text = risk.get('risk_code', '')
            row.cells[1].text = risk.get('name', '')
            row.cells[2].text = risk.get('subcategory', '')
            row.cells[3].text = str(risk.get('probability', '-'))
            row.cells[4].text = str(risk.get('impact', '-'))
            row.cells[5].text = str(risk.get('risk_score', '-'))
            row.cells[6].text = level_map.get(risk.get('risk_level', ''), '')

        doc.add_page_break()

        # 3. 风险详细分析
        doc.add_heading('三、风险详细分析', level=1)
        for i, risk in enumerate(risks):
            doc.add_heading(f'3.{i+1} {risk.get("risk_code", "")} - {risk.get("name", "")}', level=2)

            doc.add_paragraph(f'风险类别：{risk.get("category", "")} > {risk.get("subcategory", "")}')
            doc.add_paragraph(f'风险得分：{risk.get("risk_score", "")}分 '
                            f'（可能性P={risk.get("probability", "")} × 影响I={risk.get("impact", "")}）'
                            f' | 等级：{level_map.get(risk.get("risk_level", ""), "")}')
            doc.add_paragraph(f'趋势：{risk.get("trend", "")}')

            doc.add_heading('风险描述', level=3)
            doc.add_paragraph(risk.get('description', ''))

            doc.add_heading('可能后果', level=3)
            doc.add_paragraph(risk.get('possible_consequences', ''))

            doc.add_heading('当前应对措施', level=3)
            doc.add_paragraph(risk.get('current_response', ''))

            doc.add_heading('改进建议', level=3)
            doc.add_paragraph(risk.get('suggested_improvement', ''))

        doc.add_page_break()

        # 4. 预警信息
        doc.add_heading('四、预警信息', level=1)
        if warnings:
            for w in warnings:
                p = doc.add_paragraph()
                level_emoji = {"red": "🔴", "orange": "🟠", "yellow": "🟡"}
                p.add_run(f'{level_emoji.get(w.get("warning_level", ""), "")} '
                         f'{w.get("warning_message", "")}')
        else:
            doc.add_paragraph('当前无预警信息')

        # 5. 优先行动事项
        doc.add_heading('五、优先行动事项', level=1)
        for i, action in enumerate(priority_actions):
            doc.add_paragraph(
                f'{i+1}. [{action.get("risk_code", "")}] {action.get("risk_name", "")} '
                f'(得分: {action.get("risk_score", 0)})'
            )
            doc.add_paragraph(f'   行动：{action.get("action", "")}')
            doc.add_paragraph(f'   截止时间：{action.get("deadline", "")}')

        # 6. 附录
        doc.add_heading('六、附录', level=1)
        doc.add_paragraph('本报告由Agent智能风险监控管理系统自动生成。')
        doc.add_paragraph('分析框架：COSO ERM 2017')
        doc.add_paragraph('数据来源：企业公开信息（年报、公告、ESG报告等）')
        doc.add_paragraph('生成时间：' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        # 保存文件
        filename = f'{enterprise_name}_风险监控报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        filepath = os.path.join(self.REPORT_DIR, filename)
        doc.save(filepath)

        logger.info(f"Word报告已生成: {filepath}")
        return filepath

    def generate_excel_report(self,
                              enterprise_name: str,
                              risks: List[Dict],
                              stats: Dict[str, Any]) -> str:
        """生成Excel格式风险清单"""
        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment
        except ImportError:
            logger.error("openpyxl未安装")
            return ""

        wb = openpyxl.Workbook()

        # Sheet 1: 风险清单
        ws1 = wb.active
        ws1.title = "风险清单"

        headers = ['风险编号', '风险名称', '类别', '子类别', '可能性(P)',
                   '影响(I)', '风险得分', '风险等级', '趋势', '应对策略',
                   '当前措施', '改进建议']

        # 表头样式
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font = Font(color='FFFFFF', bold=True)

        for col, h in enumerate(headers, 1):
            cell = ws1.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')

        # 风险等级颜色
        level_fills = {
            'critical': PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid'),
            'high': PatternFill(start_color='FFD9B3', end_color='FFD9B3', fill_type='solid'),
            'medium': PatternFill(start_color='FFFFCC', end_color='FFFFCC', fill_type='solid'),
            'low': PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid'),
        }

        level_map = {"critical": "重大", "high": "高度关注",
                    "medium": "一般", "low": "低"}

        for row_idx, risk in enumerate(risks, 2):
            values = [
                risk.get('risk_code', ''),
                risk.get('name', ''),
                risk.get('category', ''),
                risk.get('subcategory', ''),
                risk.get('probability', ''),
                risk.get('impact', ''),
                risk.get('risk_score', ''),
                level_map.get(risk.get('risk_level', ''), ''),
                risk.get('trend', ''),
                risk.get('response_type', ''),
                risk.get('current_response', '')[:100] if risk.get('current_response') else '',
                risk.get('suggested_improvement', '')[:100] if risk.get('suggested_improvement') else '',
            ]
            for col, val in enumerate(values, 1):
                cell = ws1.cell(row=row_idx, column=col, value=val)
                cell.alignment = Alignment(wrap_text=True)
                # 给风险等级列上色
                if col == 8:
                    level = risk.get('risk_level', '')
                    if level in level_fills:
                        cell.fill = level_fills[level]

        # 调整列宽
        column_widths = [8, 25, 10, 12, 8, 8, 8, 10, 8, 10, 40, 40]
        for i, width in enumerate(column_widths, 1):
            ws1.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

        # Sheet 2: 统计概览
        ws2 = wb.create_sheet("统计概览")
        ws2.cell(row=1, column=1, value="指标").font = Font(bold=True)
        ws2.cell(row=1, column=2, value="数值").font = Font(bold=True)

        stat_data = [
            ("风险总数", stats.get('total', 0)),
            ("重大风险", stats.get('by_level', {}).get('critical', 0)),
            ("高度关注风险", stats.get('by_level', {}).get('high', 0)),
            ("一般风险", stats.get('by_level', {}).get('medium', 0)),
            ("低风险", stats.get('by_level', {}).get('low', 0)),
            ("平均风险得分", stats.get('average_score', 0)),
        ]
        for row_idx, (label, value) in enumerate(stat_data, 2):
            ws2.cell(row=row_idx, column=1, value=label)
            ws2.cell(row=row_idx, column=2, value=value)

        # 保存
        filename = f'{enterprise_name}_风险清单_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        filepath = os.path.join(self.REPORT_DIR, filename)
        wb.save(filepath)

        logger.info(f"Excel报告已生成: {filepath}")
        return filepath

    def generate_full_report_package(self,
                                     enterprise_name: str,
                                     analysis_result: Dict[str, Any]) -> Dict[str, str]:
        """
        生成完整的报告包（Word + Excel + JSON）

        Returns:
            {'word': path, 'excel': path, 'json': path}
        """
        risks = analysis_result.get('risks', [])
        stats = analysis_result.get('statistics', {})
        matrix_data = analysis_result.get('matrix_data', {})
        warnings = analysis_result.get('warnings', [])
        strategy_summary = analysis_result.get('strategy_summary', {})
        priority_actions = analysis_result.get('priority_actions', [])

        files = {}

        # Word报告
        word_path = self.generate_word_report(
            enterprise_name, risks, stats, matrix_data,
            warnings, strategy_summary, priority_actions
        )
        if word_path:
            files['word'] = word_path

        # Excel报告
        excel_path = self.generate_excel_report(enterprise_name, risks, stats)
        if excel_path:
            files['excel'] = excel_path

        # JSON数据
        json_filename = f'{enterprise_name}_风险数据_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'
        json_path = os.path.join(self.REPORT_DIR, json_filename)
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis_result, f, ensure_ascii=False, indent=2)
        files['json'] = json_path

        logger.info(f"报告包已生成: {files}")
        return files

    def generate_full_five_section_report(self, enterprise_name: str,
                                          risks: List[Dict], governance: Optional[Dict],
                                          controls: List[Dict], coverage: Dict,
                                          mechanisms: List[Dict], disclosures: List[Dict],
                                          issues: List[Dict]) -> str:
        """生成五节完整报告（匹配主报告结构）"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return ""

        doc = Document()
        title = doc.add_heading(f'{enterprise_name} 企业全面风险管理报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}\n分析引擎：Agent智能风险监控管理系统 v2.0\n理论框架：COSO ERM 2017 / ISO 31000')

        doc.add_page_break()

        # 一、摘要
        doc.add_heading('一、风险分析摘要', level=1)
        total = len(risks)
        critical = sum(1 for r in risks if r.get('risk_level') == 'critical')
        high = sum(1 for r in risks if r.get('risk_level') == 'high')
        sorted_r = sorted(risks, key=lambda r: r.get('risk_score', 0), reverse=True)
        top3 = sorted_r[:3]
        top3_text = '、'.join(r.get('name', '') for r in top3)
        avg = sum(r.get('risk_score', 0) for r in risks) / max(total, 1)

        doc.add_paragraph(
            f'本报告基于COSO ERM 2017框架和ISO 31000标准，采用Multi-Agent智能分析引擎，'
            f'对{enterprise_name}的企业公开信息进行系统化风险分析。'
            f'共识别{total}项主要风险，其中重大风险{critical}项、高度关注风险{high}项，'
            f'平均风险得分{avg:.1f}分。得分最高的三项风险为：{top3_text}。'
            f'系统自动完成风险预警检测，并基于4T策略框架（规避/降低/转移/接受）'
            f'提出了应对方案和改进建议。'
        )

        # 二、风险清单
        doc.add_heading('二、风险识别清单', level=1)
        if risks:
            table = doc.add_table(rows=1, cols=6, style='Table Grid')
            for i, h in enumerate(['编号', '风险名称', 'P', 'I', '得分', '等级']):
                table.rows[0].cells[i].text = h
            level_map = {"critical": "重大", "high": "高度关注", "medium": "一般", "low": "低"}
            for r in risks:
                row = table.add_row()
                row.cells[0].text = r.get('risk_code', '')
                row.cells[1].text = r.get('name', '')
                row.cells[2].text = str(r.get('probability', '-'))
                row.cells[3].text = str(r.get('impact', '-'))
                row.cells[4].text = str(r.get('risk_score', '-'))
                row.cells[5].text = level_map.get(r.get('risk_level', ''), '')
        doc.add_page_break()

        # 三、风险详细分析与应对
        doc.add_heading('三、风险详细分析与应对建议', level=1)
        for i, r in enumerate(risks):
            doc.add_heading(f'3.{i+1} {r.get("risk_code","")} - {r.get("name","")}', level=2)
            doc.add_paragraph(f'类别：{r.get("category","")} > {r.get("subcategory","")} | 得分：{r.get("risk_score","")} | 等级：{level_map.get(r.get("risk_level",""),"")}')
            doc.add_paragraph(f'风险描述：{r.get("description","")}')
            doc.add_paragraph(f'可能后果：{r.get("possible_consequences","")}')
            doc.add_paragraph(f'应对策略：{r.get("response_type","reduce")} - 当前措施：{r.get("current_response","")}')
            doc.add_paragraph(f'改进建议：{r.get("suggested_improvement","")}')
        doc.add_page_break()

        # 四、内控机制评估
        doc.add_heading('四、内控机制评估', level=1)
        doc.add_heading('4.1 治理结构', level=2)
        if governance:
            doc.add_paragraph(f'治理模式：{governance.get("structure_type","未设置")} | 董事会规模：{governance.get("board_size","?")}人（独董{governance.get("independent_directors","?")}人）')
            if governance.get('committees'):
                doc.add_paragraph(f'专门委员会：{governance.get("committees")}')
            if governance.get('description'):
                doc.add_paragraph(governance.get('description'))

        doc.add_heading('4.2 内控制度', level=2)
        if controls:
            for c in controls:
                doc.add_paragraph(f'{c.get("control_code","")} {c.get("name","")} [{c.get("category","")}] - {c.get("description","")[:150]} (有效性: {c.get("effectiveness","")})')
        doc.add_paragraph(f'内控覆盖率：{coverage.get("coverage_rate","0%")} | 已覆盖{coverage.get("covered_risks",0)}/{coverage.get("total_risks",0)}项风险')
        doc.add_page_break()

        # 五、监督改进
        doc.add_heading('五、监督与改进', level=1)
        doc.add_heading('5.1 监督机制', level=2)
        if mechanisms:
            for m in mechanisms:
                doc.add_paragraph(f'{m.get("mechanism_name","")} [{m.get("mechanism_type","")}] - 频率：{m.get("frequency","")} - {m.get("description","")[:200]}')
        else:
            doc.add_paragraph('暂未记录监督机制信息')

        doc.add_heading('5.2 信息披露', level=2)
        if disclosures:
            for d in disclosures:
                doc.add_paragraph(f'{d.get("disclosure_type","")} - {d.get("title","")} (发布日期：{d.get("publish_date","")[:10] if d.get("publish_date") else "?"})')
        else:
            doc.add_paragraph('暂未记录信息披露')

        doc.add_heading('5.3 发现问题与改进', level=2)
        if issues:
            for iss in issues:
                doc.add_paragraph(f'{iss.get("issue_code","")} {iss.get("title","")} [{iss.get("severity","")}] - 状态：{iss.get("status","")} - {iss.get("proposed_action","")[:200]}')
        else:
            doc.add_paragraph('暂无问题记录')

        doc.add_heading('5.4 持续改进建议', level=2)
        doc.add_paragraph('基于上述分析，建议企业：①建立健全风险管理委员会；②定期开展全面风险评估；③加强内控制度的执行监督；④建立风险管理的PDCA持续改进循环。')

        # 保存
        filename = f'{enterprise_name}_全面风险管理报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        filepath = os.path.join(self.REPORT_DIR, filename)
        doc.save(filepath)
        logger.info(f"五节完整报告已生成: {filepath}")
        return filepath
