# -*- coding: utf-8 -*-
"""
数据接入模块 - 智能解析企业公开信息
支持：年报/ESG报告DOCX解析、结构化风险数据提取、风险披露章节识别
"""

import re
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime

logger = logging.getLogger(__name__)


class DataIngestionEngine:
    """
    企业公开信息智能接入引擎

    功能：
    1. 解析上传的DOCX/TXT文件，提取风险相关章节
    2. 从文本中智能提取结构化风险数据
    3. 识别风险类别、描述、后果等关键字段
    4. 生成符合主报告格式的结构化风险条目
    """

    # 风险章节识别关键词（中英文）
    RISK_SECTION_MARKERS = [
        # 中文年报常见标题
        '风险因素', '风险管理', '风险披露', '重大风险',
        '风险分析', '风险与不确定', '风险提示', '经营风险',
        '内部控制', '公司治理', '风险管理体系',
        # 英文年报常见标题
        'risk factors', 'risk management', 'risk disclosure',
        'principal risks', 'key risks', 'material risks',
        'risk and uncertainty', 'risk assessment',
        # ESG报告
        'ESG', '环境、社会及管治', '可持续发展', '碳中和',
        '气候变化风险', '数据安全', '隐私保护',
    ]

    # 行业特定风险关键词
    INDUSTRY_RISK_PATTERNS = {
        '科技/互联网': [
            (r'监管|政策|合规|牌照|许可|审批|版号|反垄断|算法|数据安全|隐私|个人信息|网络安全',
             '监管政策风险'),
            (r'竞争|市场份|对手|字节|阿里|腾讯|美团|拼多多|用户时|流量',
             '行业竞争风险'),
            (r'制裁|实体清单|出口管制|贸易战|芯片|GPU|脱钩|地缘|美国|关税',
             '地缘政治风险'),
            (r'AI|人工智|大模型|算法|GPT|技术迭代|颠覆|创新',
             '技术变革风险'),
            (r'数据泄露|信息安全|黑客|攻击|系统故障|宕机|容灾',
             '数据安全风险'),
        ],
        '制造业': [
            (r'供应链|原材料|价格波动|短缺|中断|物流|库存',
             '供应链风险'),
            (r'产能|开工率|利用率|过剩|产线|良率|品控|质量',
             '产能与质量风险'),
            (r'技术迭代|新能|固态|锂电|氢能|颠覆|替代|下一代',
             '技术变革风险'),
            (r'关税|贸易|反倾销|调查|制裁|出口|进口|地缘',
             '国际贸易风险'),
            (r'排放|碳|ESG|环保|能耗|双碳|绿色|可持续',
             'ESG合规风险'),
            (r'劳工|用工|薪酬|罢工|劳动|人才|技工',
             '人力资源风险'),
        ],
        '金融': [
            (r'信用|违约|坏账|不良资产|逾期|拨备',
             '信用风险'),
            (r'市场|利率|汇率|波动|利差|净息差|估值',
             '市场风险'),
            (r'流动|资金|存款|兑付|挤兑|备付',
             '流动性风险'),
            (r'操作|内控|合规|反洗钱|KYC|AML|欺诈|案件',
             '操作风险'),
            (r'监管|政策|资本充|巴塞尔|宏观审慎|银保监|央行',
             '监管合规风险'),
            (r'金融科技|互联网金|数字货|区块链|支付|第三方',
             '金融科技风险'),
        ],
        '消费': [
            (r'消费|需求|购买力|降级|复苏|疲软|下行',
             '消费需求风险'),
            (r'品牌|声誉|口碑|舆情|负面|投诉|315',
             '品牌声誉风险'),
            (r'渠道|电商|线上|线下|经销商|终端|分销',
             '渠道风险'),
            (r'食品|安全|质量|卫生|添加剂|召回|检测',
             '食品安全风险'),
            (r'竞争|价格战|促销|营销|获客|留存|复购',
             '市场竞争风险'),
        ],
        '能源': [
            (r'油价|天然气|煤|碳|国家管|价改|定价',
             '价格波动风险'),
            (r'安全|事故|爆炸|泄漏|井喷|矿难|HSE',
             '安全生产风险'),
            (r'碳排|碳中和|能源转|新能|光伏|风电|氢',
             '能源转型风险'),
            (r'地缘|制裁|中东|OPEC|俄罗斯|管道|海峡',
             '地缘政治风险'),
        ],
    }

    # 财务指标风险阈值
    FINANCIAL_RISK_THRESHOLDS = {
        '资产负债率': {'high': 0.70, 'medium': 0.50, 'unit': '%'},
        '流动比率': {'low': 1.0, 'medium': 1.5, 'unit': ''},
        '应收账款周转天数': {'high': 120, 'medium': 60, 'unit': '天'},
        '研发费用率': {'low': 0.03, 'medium': 0.05, 'unit': '%'},
        '前5大客户集中度': {'high': 0.50, 'medium': 0.30, 'unit': '%'},
    }

    def __init__(self):
        self.parsed_documents: List[Dict] = []
        self.extracted_risks: List[Dict] = []

    def parse_document(self, filepath: str) -> Dict[str, Any]:
        """
        解析上传的文档（DOCX/TXT），提取风险相关内容

        Returns:
            {
                'filepath': str,
                'full_text': str,
                'risk_sections': [{'title': str, 'content': str}],
                'financial_data': {'营收': xxx, '净利润': xxx, ...},
                'enterprise_name': str or None,
                'risk_count': int,
                'word_count': int
            }
        """
        logger.info(f"解析文档: {filepath}")

        # 提取文本
        full_text = self._extract_text(filepath)
        if not full_text:
            return {'error': '无法提取文本内容', 'filepath': filepath}

        # 识别企业名称
        enterprise_name = self._extract_enterprise_name(full_text)

        # 提取风险相关章节
        risk_sections = self._extract_risk_sections(full_text)

        # 提取财务数据
        financial_data = self._extract_financial_data(full_text)

        # 识别风险条目
        extracted_risks = self._extract_structured_risks(full_text, risk_sections)

        result = {
            'filepath': filepath,
            'filename': os.path.basename(filepath) if '/' in filepath or '\\' in filepath else filepath,
            'full_text': full_text[:50000],  # 截断避免过大
            'full_text_length': len(full_text),
            'risk_sections': risk_sections,
            'financial_data': financial_data,
            'enterprise_name': enterprise_name,
            'extracted_risks': extracted_risks,
            'risk_count': len(extracted_risks),
            'word_count': len(full_text),
            'parsed_at': datetime.utcnow().isoformat()
        }

        self.parsed_documents.append(result)
        self.extracted_risks.extend(extracted_risks)

        logger.info(f"文档解析完成: {len(risk_sections)} 个风险章节, "
                   f"{len(extracted_risks)} 条风险, {len(full_text)} 字符")
        return result

    def parse_text_block(self, text: str, industry: str = '') -> Dict[str, Any]:
        """
        解析用户粘贴的文本块（年报摘要、新闻等）
        """
        risk_sections = self._extract_risk_sections(text)
        extracted_risks = self._extract_structured_risks(text, risk_sections, industry)
        financial_data = self._extract_financial_data(text)

        return {
            'risk_sections': risk_sections if risk_sections else [{'title': '全文', 'content': text[:2000]}],
            'financial_data': financial_data,
            'extracted_risks': extracted_risks,
            'risk_count': len(extracted_risks),
            'word_count': len(text),
        }

    def generate_enterprise_summary(self, parsed_result: Dict) -> str:
        """基于解析结果生成企业风险概况摘要"""
        risks = parsed_result.get('extracted_risks', [])
        sections = parsed_result.get('risk_sections', [])

        if not risks:
            return "未识别到明确的风险条目，请提供更详细的企业公开信息。"

        parts = []
        parts.append(f"从文档中识别出 {len(risks)} 项潜在风险,"
                    f"涉及 {len(sections)} 个风险相关章节。")

        # 按类别统计
        categories = {}
        for r in risks:
            cat = r.get('subcategory', r.get('category', '未分类'))
            categories[cat] = categories.get(cat, 0) + 1

        cat_summary = "；".join([f"{k}：{v}项" for k, v in categories.items()])
        parts.append(f"风险类别分布：{cat_summary}。")

        # 财务指标
        fin = parsed_result.get('financial_data', {})
        if fin:
            fin_items = []
            if fin.get('营收'):
                fin_items.append(f"营收{fin['营收']}")
            if fin.get('净利润'):
                fin_items.append(f"净利润{fin['净利润']}")
            if fin.get('资产负债率'):
                fin_items.append(f"资产负债率{fin['资产负债率']}")
            if fin_items:
                parts.append(f"关键财务指标：{'，'.join(fin_items)}。")

        return '\n'.join(parts)

    # ========== 内部方法 ==========

    def _extract_text(self, filepath: str) -> str:
        """从文件提取文本"""
        ext = filepath.lower().rsplit('.', 1)[-1] if '.' in filepath else ''

        if ext in ['docx', 'doc']:
            try:
                from docx import Document
                doc = Document(filepath)
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p.text.strip())
                # 也提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(
                            cell.text.strip() for cell in row.cells
                            if cell.text.strip()
                        )
                        if row_text:
                            paragraphs.append(row_text)
                return '\n'.join(paragraphs)
            except Exception as e:
                logger.error(f"DOCX解析失败: {e}")
                return ""

        elif ext in ['txt', '']:
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        return f.read()
                except (UnicodeDecodeError, UnicodeError):
                    continue

        return ""

    def _extract_enterprise_name(self, text: str) -> Optional[str]:
        """从文本中提取企业名称"""
        # 常见模式
        patterns = [
            r'([一-龥]{2,8}(?:股份|集团|控股|科技|银行|保险|证券|基金|实业|有限|汽车|医药|能源|地产|建设)有限公司)',
            r'([一-龥]{2,8}(?:股份|集团|控股|科技|银行|保险|证券|基金|实业|有限|汽车|医药|能源|地产|建设)(?:公司|企业))',
        ]

        for pattern in patterns:
            match = re.search(pattern, text[:2000])
            if match:
                return match.group(1)

        # 搜索常见企业名称关键词
        for line in text.split('\n')[:50]:
            line = line.strip()
            if any(kw in line for kw in ['公司名称', '企业名称', '股票简称', '证券简称']):
                parts = line.replace('：', ':').split(':')
                if len(parts) >= 2:
                    name = parts[-1].strip()
                    if 4 <= len(name) <= 20:
                        return name

        return None

    def _extract_risk_sections(self, text: str) -> List[Dict[str, str]]:
        """
        提取风险相关章节
        Returns: [{'title': str, 'content': str, 'line_start': int, 'line_end': int}]
        """
        lines = text.split('\n')
        sections = []
        current_section = None

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue

            # 检查是否是新章节标题
            is_section_title = False
            for marker in self.RISK_SECTION_MARKERS:
                if marker in stripped and len(stripped) < 60:
                    is_section_title = True
                    break

            if is_section_title:
                if current_section and current_section['content'].strip():
                    sections.append(current_section)
                current_section = {
                    'title': stripped,
                    'content': '',
                    'line_start': i,
                    'line_end': i
                }
            elif current_section:
                current_section['content'] += line + '\n'
                current_section['line_end'] = i

        # 添加最后一个章节
        if current_section and current_section['content'].strip():
            sections.append(current_section)

        # 如果没有找到明确的风险章节，尝试寻找包含大量风险关键词的段落
        if not sections:
            risk_keywords_count = sum(
                1 for line in lines
                if any(kw in line for kw in ['风险', 'risk', '不确定', '挑战'])
            )
            if risk_keywords_count > 0:
                sections.append({
                    'title': '全文（已自动识别风险相关内容）',
                    'content': '\n'.join(lines),
                    'line_start': 0,
                    'line_end': len(lines)
                })

        return sections

    def _extract_structured_risks(self, text: str,
                                   risk_sections: List[Dict],
                                   industry: str = '') -> List[Dict[str, Any]]:
        """
        从文本中智能提取结构化风险条目
        结合风险章节内容和行业关键词进行识别
        """
        risks = []
        risk_code_counter = 1

        # 选择行业匹配模式
        matched_industry = self._match_industry(text, industry)
        patterns = self.INDUSTRY_RISK_PATTERNS.get(
            matched_industry,
            self.INDUSTRY_RISK_PATTERNS['科技/互联网']
        )

        # 合并所有风险章节内容
        all_risk_text = '\n'.join(
            s['content'] for s in risk_sections
        ) if risk_sections else text

        # 对每种风险类别进行匹配
        for pattern, category_name in patterns:
            matches = re.finditer(pattern, all_risk_text, re.IGNORECASE)
            matched_sentences = []

            for m in matches:
                # 提取包含匹配的完整句子
                start = max(0, m.start() - 100)
                end = min(len(all_risk_text), m.end() + 200)
                context = all_risk_text[start:end]

                # 截取完整句子
                sentence = self._extract_sentence(context, m.group())
                if sentence and len(sentence) > 15:
                    matched_sentences.append(sentence)
                if len(matched_sentences) >= 3:
                    break

            if not matched_sentences:
                continue  # 此类别无匹配

            # 合并匹配的句子作为描述
            description = '。'.join(matched_sentences[:3])

            # 生成可能后果
            consequences = self._generate_consequences(category_name)

            # 生成风险名称
            risk_name = self._generate_risk_name(category_name, matched_industry)

            # 评估P和I
            prob, impact = self._estimate_pi(description, category_name)

            risk = {
                'risk_code': f'R{risk_code_counter}',
                'name': risk_name,
                'category': '外部风险' if category_name in [
                    '监管政策风险', '行业竞争风险', '地缘政治风险',
                    '技术变革风险', '国际贸易风险', '消费需求风险',
                    '市场价格风险', '金融科技风险', '能源转型风险'
                ] else '内部风险',
                'subcategory': category_name,
                'description': description[:800],
                'possible_consequences': consequences,
                'information_source': '企业提供的公开信息文本',
                'probability': prob,
                'impact': impact,
                'status': 'active',
                'trend': 'stable',
                'ai_identified': True,
                'ai_assessment_notes': f'从文本中自动提取（匹配{matched_industry}行业模式）',
            }
            risk['risk_score'] = prob * impact
            risk['risk_level'] = (
                'critical' if risk['risk_score'] >= 15 else
                'high' if risk['risk_score'] >= 10 else
                'medium' if risk['risk_score'] >= 5 else 'low'
            )
            risks.append(risk)
            risk_code_counter += 1

        # 如果匹配到的风险太少，尝试用通用模式
        if len(risks) < 3:
            generic_risks = self._extract_generic_risks(text, risk_code_counter)
            risks.extend(generic_risks)

        return risks[:12]  # 最多12个

    def _match_industry(self, text: str, given_industry: str) -> str:
        """匹配行业"""
        if given_industry:
            for key in self.INDUSTRY_RISK_PATTERNS:
                if key in given_industry or given_industry in key:
                    return key

        # 基于文本关键词推断
        industry_keywords = {
            '科技/互联网': ['互联网', '软件', '科技', '平台', 'APP', 'SaaS', '云', '数据'],
            '制造业': ['制造', '工厂', '产能', '产线', '电池', '汽车', '芯片', '半导体'],
            '金融': ['银行', '金融', '保险', '证券', '基金', '贷款', '存款', '理财'],
            '消费': ['零售', '消费', '食品', '饮料', '白酒', '乳业', '品牌'],
            '能源': ['能源', '石油', '天然气', '电力', '煤炭', '光伏', '风电'],
        }

        text_2000 = text[:2000].lower()
        for industry, keywords in industry_keywords.items():
            score = sum(1 for kw in keywords if kw in text_2000)
            if score >= 2:
                return industry

        return '科技/互联网'  # 默认

    def _extract_sentence(self, context: str, keyword: str) -> str:
        """从上下文中提取包含关键词的完整句子"""
        # 按句号、分号、换行分割
        sentences = re.split(r'[。；\n]', context)
        for s in sentences:
            if keyword.lower() in s.lower() and len(s.strip()) > 10:
                return s.strip()
        return context.strip()[:200]

    def _generate_risk_name(self, category: str, industry: str) -> str:
        """生成风险名称"""
        templates = {
            '监管政策风险': f'{industry}行业监管政策变化带来的合规风险',
            '行业竞争风险': f'{industry}行业竞争加剧与市场份额风险',
            '地缘政治风险': '国际地缘政治环境变化风险',
            '技术变革风险': '技术迭代与创新风险',
            '供应链风险': '供应链稳定性风险',
            '产能与质量风险': '产能利用与产品质量风险',
            '国际贸易风险': '国际贸易环境变化风险',
            'ESG合规风险': 'ESG与可持续发展合规风险',
            '人力资源风险': '人才吸引与保留风险',
            '信用风险': '信用与资产质量风险',
            '市场风险': '市场价格波动风险',
            '流动性风险': '流动性管理风险',
            '操作风险': '操作与内控合规风险',
            '监管合规风险': '金融监管政策合规风险',
            '金融科技风险': '金融科技与数字化转型风险',
            '消费需求风险': '消费需求变化风险',
            '品牌声誉风险': '品牌声誉管理风险',
            '渠道风险': '渠道管理风险',
            '食品安全风险': '食品质量安全风险',
            '市场竞争风险': '市场竞争与获客风险',
            '价格波动风险': '能源价格波动风险',
            '安全生产风险': '安全生产与HSE风险',
            '能源转型风险': '能源结构转型风险',
            '数据安全风险': '数据安全与隐私保护风险',
        }
        return templates.get(category, f'{category}')

    def _generate_consequences(self, category: str) -> str:
        """生成可能后果描述"""
        consequences_map = {
            '监管政策风险': '监管处罚导致业务暂停或高额罚款；合规成本大幅增加；产品准入受限影响市场拓展',
            '行业竞争风险': '市场份额被侵蚀；价格竞争导致利润率下降；竞争优势被削弱；用户流失',
            '地缘政治风险': '海外市场准入受限；关税成本大幅增加；关键零部件和技术供应中断',
            '技术变革风险': '技术路线被颠覆导致前期投入减值；产品竞争力下降；研发回报不及预期',
            '供应链风险': '原材料短缺导致生产停滞；价格波动影响成本和毛利；交付延迟损害客户关系',
            '产能与质量风险': '产能过剩导致资产减值；产品召回损害品牌；良率问题增加成本',
            '国际贸易风险': '出口受阻影响海外收入；汇率波动侵蚀利润；贸易壁垒增加',
            'ESG合规风险': '碳排放超标面临罚款；ESG评级下降影响融资和品牌；环保诉讼风险',
            '人力资源风险': '核心人才流失影响研发和运营；用工成本上升；劳资纠纷损害声誉',
            '信用风险': '坏账增加影响利润；资产质量恶化；拨备不足',
            '市场风险': '利率/汇率波动影响投资收益；证券估值缩水；息差收窄',
            '流动性风险': '流动性不足影响正常经营；融资成本上升；触发监管红线',
            '操作风险': '内控失效导致损失；欺诈事件；系统故障影响服务连续性',
            '数据安全风险': '数据泄露导致用户信任危机和监管重罚；被攻击导致业务中断',
        }
        return consequences_map.get(category, '对企业经营业绩和财务状况产生不利影响')

    def _estimate_pi(self, description: str, category: str) -> Tuple[int, int]:
        """基于描述内容估计P和I值"""
        desc_lower = description.lower()

        # 估计可能性
        prob = 3
        if any(kw in desc_lower for kw in ['持续', '频繁', '已经', '已发生', '严峻', '激烈', '恶化']):
            prob = 4
        elif any(kw in desc_lower for kw in ['可能', '潜在', '远期', '尚未', '如果', '预计']):
            prob = 2

        # 估计影响
        impact = 3
        # 高影响类别
        high_impact_cats = ['监管政策风险', '地缘政治风险', '安全生产风险', '技术变革风险', '信用风险']
        # 中等影响类别
        medium_impact_cats = ['行业竞争风险', '供应链风险', 'ESG合规风险', '人力资源风险', '数据安全风险']

        if category in high_impact_cats:
            impact = 4
        elif category in medium_impact_cats:
            impact = 3
        else:
            impact = 2

        # 描述中包含严重性指标
        if any(kw in desc_lower for kw in ['重大', '严重', '核心', '关键', '威胁生存']):
            impact = min(5, impact + 1)

        return prob, impact

    def _extract_financial_data(self, text: str) -> Dict[str, str]:
        """从文本中提取财务指标"""
        financial = {}

        # 营收
        rev_patterns = [
            r'营[业收]*[入收].*?([\d,]+\.?\d*)\s*[万亿]?元',
            r'(?:revenue|Revenue).*?([\d,]+\.?\d*)',
        ]
        for pattern in rev_patterns:
            match = re.search(pattern, text[:3000])
            if match:
                financial['营收'] = match.group(0)[:50]
                break

        # 净利润
        profit_patterns = [
            r'净利[润利].*?([\d,]+\.?\d*)\s*[万亿]?元',
            r'归属于.*?净利[润利].*?([\d,]+\.?\d*)',
        ]
        for pattern in profit_patterns:
            match = re.search(pattern, text[:3000])
            if match:
                financial['净利润'] = match.group(0)[:50]
                break

        # 资产负债率
        debt_match = re.search(
            r'资产负债[率比].*?(\d+\.?\d*)\s*%',
            text[:5000]
        )
        if debt_match:
            financial['资产负债率'] = debt_match.group(0)[:30]

        # 增长率
        growth_match = re.search(
            r'(?:同比|较上年|compared).*?(?:增[长加]|下降|减少).*?(\d+\.?\d*)\s*%',
            text[:3000]
        )
        if growth_match:
            financial['增长率'] = growth_match.group(0)[:40]

        return financial

    def _extract_generic_risks(self, text: str, start_code: int = 1) -> List[Dict]:
        """通用风险提取（当行业匹配不足时的兜底）"""
        risks = []
        # 全局风险关键词扫描
        global_patterns = [
            (r'(?:监管|政策|合规|牌照|许可|审批|法规)', '监管政策风险'),
            (r'(?:竞争|市场份额|价格战|内卷)', '行业竞争风险'),
            (r'(?:人才|招聘|用工|薪酬|流失|裁员)', '人力资源风险'),
            (r'(?:财务|资金|负债|现金流|应收|应付|借款)', '财务管理风险'),
            (r'(?:数据|隐私|信息|安全|漏洞|攻击|泄露)', '数据安全风险'),
            (r'(?:技术|研发|创新|迭代|专利|知识产权)', '技术风险'),
            (r'(?:供应链|采购|原材料|供应商|交付)', '运营风险'),
        ]

        text_5000 = text[:5000]
        for pattern, cat_name in global_patterns:
            if re.search(pattern, text_5000, re.IGNORECASE):
                risk = {
                    'risk_code': f'R{start_code}',
                    'name': cat_name,
                    'category': '外部风险' if cat_name in ['监管政策风险', '行业竞争风险'] else '内部风险',
                    'subcategory': cat_name,
                    'description': f'基于企业提供的公开信息，识别出{cat_name}。'
                                  f'该风险可能对企业经营产生影响，需要进一步详细分析。',
                    'possible_consequences': '对企业经营业绩产生不利影响',
                    'information_source': '企业提供的公开信息文本',
                    'probability': 3,
                    'impact': 3,
                    'status': 'active',
                    'trend': 'stable',
                    'ai_identified': True,
                }
                risk['risk_score'] = 9
                risk['risk_level'] = 'medium'
                risks.append(risk)
                start_code += 1

        return risks


import os
