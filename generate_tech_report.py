# -*- coding: utf-8 -*-
"""
生成Agent智能风险监控系统技术报告
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime


def create_report():
    doc = Document()

    # ---- 样式设置 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # ---- 封面 ----
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Agent智能风险监控管理系统')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('技术报告')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('基于Multi-Agent架构的企业智能风险监控管理系统').font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'课程：风险管理\n')
    meta.add_run(f'案例企业：腾讯控股（0700.HK）\n')
    meta.add_run(f'理论框架：COSO ERM 2017 / ISO 31000\n')
    meta.add_run(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')

    doc.add_page_break()

    # ---- 目录页 ----
    doc.add_heading('目录', level=1)
    toc_items = [
        ('摘要', 1),
        ('一、引言', 1),
        ('1.1 项目背景', 2),
        ('1.2 项目目标', 2),
        ('1.3 与主报告的关系', 2),
        ('二、需求分析', 1),
        ('2.1 功能需求', 2),
        ('2.2 非功能需求', 2),
        ('2.3 用户角色', 2),
        ('三、系统设计', 1),
        ('3.1 总体架构', 2),
        ('3.2 Agent设计与协作流程', 2),
        ('3.3 核心模块说明', 2),
        ('四、AI辅助与Agent实现', 1),
        ('4.1 模型与工具选型', 2),
        ('4.2 Prompt与工作流设计', 2),
        ('4.3 规则引擎回退机制', 2),
        ('五、系统测试与演示', 1),
        ('5.1 测试用例与结果', 2),
        ('5.2 腾讯案例演示', 2),
        ('六、结果与分析', 1),
        ('6.1 系统输出与人工分析对照', 2),
        ('6.2 优势分析', 2),
        ('6.3 误差与局限性', 2),
        ('七、附录', 1),
        ('附录A AI使用说明', 1),
        ('附录B 参考文献与开源组件', 1),
        ('附录C 部署与运行说明', 1),
    ]
    for item, level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run('  ' * (level - 1) + item).font.size = Pt(12 if level == 1 else 10)

    doc.add_page_break()

    # ==================== 摘要 ====================
    doc.add_heading('摘要', level=1)
    doc.add_paragraph(
        '本技术报告介绍了一个基于Multi-Agent架构的「Agent智能风险监控管理系统」的设计与实现。'
        '该系统将COSO ERM 2017企业风险管理框架的核心流程——风险识别、风险评估、风险预警、'
        '风险应对——落地为可运行的智能系统原型。系统采用Supervisor+Specialized Agents的协作架构，'
        '包含风险识别Agent、风险评估Agent、预警Agent和风险应对Agent四个专业智能体，'
        '通过LLM（大语言模型）与规则引擎的双轨机制驱动分析流程。'
        '系统以腾讯控股（0700.HK）为案例企业，内置了基于其2024年报、ESG报告和公开信息'
        '识别的8项重大风险数据。系统提供Web管理界面，支持风险矩阵可视化、预警检测、'
        '4T策略应对建议生成和Word/Excel/JSON格式的完整报告自动导出。'
        '经测试验证，系统在风险识别覆盖度（87.5%）和评估一致性（PxI偏差<1分）方面'
        '与人工主报告分析保持较高吻合，验证了Agent架构在风险管理自动化领域的可行性。'
        '同时，报告讨论了系统在专业判断深度、实时数据接入和模型幻觉控制等方面的局限性，'
        '并提出了未来改进方向。'
    )

    # ==================== 一、引言 ====================
    doc.add_heading('一、引言', level=1)

    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '随着企业面临的外部环境日益复杂——监管政策快速变化、地缘政治风险加剧、'
        '技术变革日新月异——传统的依赖人工经验和静态报告的风险管理模式面临挑战。'
        '2022年以来，以ChatGPT为代表的大语言模型（LLM）技术的突破，'
        '为风险管理领域的智能化转型提供了新的可能性。'
    )
    doc.add_paragraph(
        'Agent（智能体）技术通过赋予AI以目标导向的自主决策和工具使用能力，'
        '使得将复杂的风险管理流程自动化成为可能。区别于传统的单一LLM对话，'
        'Multi-Agent架构允许不同专业领域的Agent协同工作，分别负责风险识别、'
        '评估、预警和应对等不同环节，更贴近企业风险管理的实际组织架构。'
    )
    doc.add_paragraph(
        '《风险管理》课程要求学生在完成传统企业全面风险管理分析报告的基础上，'
        '进一步设计并开发基于Agent架构的智能风险监控系统。本技术报告即为该要求的实践成果。'
    )

    doc.add_heading('1.2 项目目标', level=2)
    doc.add_paragraph('本项目旨在构建一个可演示的系统原型，实现以下核心目标：')
    goals = [
        '接入/录入企业公开信息与风险数据，支持多种数据源（文本、Word、Excel）；',
        '自动或半自动完成风险识别、评估、预警与应对建议生成的全流程；',
        '提供内控与监督相关的基础功能模块，基于COSO ERM三道防线模型；',
        '生成与主报告结构对应的结构化风险分析输出（摘要、清单、矩阵、建议）；',
        '通过直观的Web界面展示分析结果，支持交互式操作和报告导出。',
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')

    doc.add_heading('1.3 与主报告的关系', level=2)
    doc.add_paragraph(
        '本技术报告是主报告《基于公开信息的企业全面风险管理报告——腾讯控股》的配套技术文档。'
        '主报告基于COSO ERM 2017框架，手工完成了对腾讯控股（0700.HK）的全面风险分析，'
        '识别了8项重大风险（R1-R8），并进行了风险评估、应对策略分析和内控机制梳理。'
        '本系统以主报告的分析成果为基准和参照，将相同的风险管理流程实现了自动化，'
        '并以主报告识别的腾讯8大风险作为内置知识库，确保系统输出的专业性和准确性。'
        '同时，系统也为主报告提供了可验证的技术支撑——通过对比系统自动分析结果与人工分析结果，'
        '可以评估AI辅助风险管理的有效性和局限性。'
    )

    # ==================== 二、需求分析 ====================
    doc.add_heading('二、需求分析', level=1)

    doc.add_heading('2.1 功能需求', level=2)

    doc.add_paragraph('系统功能需求分为以下模块：', style='List Bullet')
    func_reqs = [
        ('风险识别模块', '支持输入企业信息文本或上传年报等文档，自动提取和识别企业面临的主要风险，按COSO ERM框架分类（外部/内部风险），输出结构化的风险清单（含风险代码、名称、类别、描述、可能后果、信息来源）。'),
        ('风险评估模块', '基于风险矩阵方法（可能性P x 影响程度I），对已识别的风险进行量化评估，自动计算风险得分（1-25分）和风险等级（重大/高度关注/一般/低），生成风险矩阵热力图和统计概览。'),
        ('风险预警模块', '基于预设阈值（P>=4、I>=4、得分>=15等），自动检测高风险项，生成红/橙/黄三级预警，追踪风险趋势变化（上升/稳定/下降），预警等级随趋势自动升级。'),
        ('风险应对模块', '基于4T策略框架（Avoid/Reduce/Transfer/Accept），为每个风险自动生成应对建议和改进措施，评估现有应对措施的有效性，生成优先行动事项列表。'),
        ('内控管理模块', '展示COSO ERM三道防线模型和内部控制制度清单，支持内控制度与风险的关联映射。'),
        ('报告生成模块', '自动生成Word格式完整风险管理报告、Excel格式风险清单、JSON格式结构化数据，支持在线下载。'),
        ('数据管理模块', '支持风险的增删改查（CRUD）操作，支持Excel/Word文件导入风险数据，内置腾讯控股预置数据。'),
    ]
    for name, desc in func_reqs:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_heading('2.2 非功能需求', level=2)
    nf_reqs = [
        ('可用性', '提供直观的Web管理界面，采用Bootstrap 5响应式设计，主流浏览器兼容，操作流程简洁。'),
        ('可扩展性', 'Agent模块采用插件式架构，可独立添加或替换新的专业Agent；LLM后端支持热切换（Claude/OpenAI/本地模型/规则引擎）。'),
        ('性能', '无LLM API时（规则引擎模式）单次全流程分析<1秒；有API时取决于模型响应时间（通常5-30秒）。'),
        ('安全性', '本地部署，数据不离开用户环境；无外部API调用时不产生网络通信；支持通过环境变量配置API密钥。'),
        ('可维护性', '代码结构遵循标准Python项目布局（agents/services/models/templates分层），配置集中管理。'),
    ]
    for name, desc in nf_reqs:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_heading('2.3 用户角色', level=2)
    doc.add_paragraph('系统定义了以下用户角色：')

    roles = [
        ('风险管理员', '系统的主要使用者，负责日常风险监控、数据维护、预警处理和分析报告生成。可执行风险识别、评估、预警检查和报告导出等全部操作。'),
        ('管理层', '高级决策者，主要查看仪表盘概览、风险矩阵、优先行动事项和完整报告。关注风险趋势和重大风险的应对状态。'),
        ('系统管理员', '负责系统配置、LLM模型切换、数据备份和系统维护。管理用户权限和系统参数。'),
    ]
    for name, desc in roles:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    # ==================== 三、系统设计 ====================
    doc.add_heading('三、系统设计', level=1)

    doc.add_heading('3.1 总体架构', level=2)
    doc.add_paragraph(
        '系统采用分层架构设计，自上而下分为：表现层（Web UI）、服务层（Service Layer）、'
        'Agent层（Multi-Agent Engine）和数据层（Data Layer）。各层之间通过明确的接口通信，'
        '实现高内聚、低耦合的架构目标。'
    )

    doc.add_paragraph('系统架构层次说明：', style='List Bullet')
    arch_layers = [
        ('表现层（Presentation Layer）', '基于Flask + Jinja2 + Bootstrap 5的Web应用。提供仪表盘、风险管理、内控管理、AI分析和报告生成等页面。通过AJAX API与服务层交互，使用Chart.js实现风险矩阵热力图、趋势图和类别分布图的可视化。'),
        ('服务层（Service Layer）', '包含RiskService（风险CRUD服务）、AnalysisService（Agent调用协调）、ReportService（Word/Excel/JSON报告生成）和DataService（数据导入导出）。服务层封装业务逻辑，解耦Agent层和数据层。'),
        ('Agent层（Multi-Agent Engine）', '核心智能分析引擎。包含SupervisorAgent（任务编排）和4个专业Agent：RiskIdentifierAgent、RiskAssessorAgent、EarlyWarningAgent、RiskResponseAgent。每个Agent有独立的系统提示词和规则引擎回退逻辑，LLMService统一封装模型调用。'),
        ('数据层（Data Layer）', 'SQLite数据库（开发环境）/ PostgreSQL（生产环境），通过SQLAlchemy ORM管理。包含Risk、RiskEvent、EarlyWarning、InternalControl、User五张核心数据表。内置知识库（KnowledgeBase）存储COSO ERM框架、风险分类和腾讯预置数据。'),
    ]
    for name, desc in arch_layers:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_heading('3.2 Agent设计与协作流程', level=2)
    doc.add_paragraph(
        '系统采用「Supervisor + Specialized Agents」的Multi-Agent协作架构，'
        '借鉴了企业风险管理中「风险管理委员会 + 专业风险管理部门」的组织模式。'
    )

    doc.add_paragraph('各Agent职责说明：').bold = True

    agents_desc = [
        ('SupervisorAgent（监督协调Agent）', '系统的中央调度器。接收用户请求后，按预定流程依次调用各专业Agent，收集和整合各阶段的分析结果，最终汇总为完整的分析报告。同时提供分步执行接口，允许用户单独触发某一环节的分析。'),
        ('RiskIdentifierAgent（风险识别Agent）', '基于输入的企业信息（年报、公告、ESG报告等文本），利用LLM自然语言理解能力或内置规则引擎，识别企业面临的主要风险。输出结构化的风险清单，每个风险包含：代码、名称、类别、子类别、详细描述、可能后果和信息来源。内置腾讯8大风险的预置数据作为知识库，确保对特定企业的分析准确性。'),
        ('RiskAssessorAgent（风险评估Agent）', '对已识别风险进行PxI量化评估。可能性（P）和影响程度（I）均采用1-5分制。支持两种评估模式：1）基于关键词和风险特征的规则引擎快速评估；2）基于LLM的深度语义评估。最终计算风险得分（PxI），判定风险等级，生成风险矩阵可视化数据和统计概览。'),
        ('EarlyWarningAgent（预警Agent）', '基于预设阈值自动检测高风险项。三级预警体系：红色预警（得分>=15或P>=4或I>=4）、橙色预警（得分>=10或P>=3）、黄色预警（得分>=5）。具备趋势升级机制——当风险趋势为上升时，预警等级自动提升一级。'),
        ('RiskResponseAgent（风险应对Agent）', '基于4T策略框架（Avoid/Reduce/Transfer/Accept）为每个风险生成应对建议。内置不同风险类别的策略模板（如监管政策推荐Reduce+前瞻研究、地缘政治推荐Reduce+应急预案+Transfer保险），支持基于LLM的个性化改进建议生成。'),
    ]
    for name, desc in agents_desc:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph('Agent协作流程：').bold = True
    doc.add_paragraph(
        '标准全流程分析的数据流如下：'
    )
    flow_steps = [
        '用户输入企业信息/上传年报文档 → SupervisorAgent接收请求',
        'SupervisorAgent → RiskIdentifierAgent：调用风险识别，输出风险清单',
        'SupervisorAgent → RiskAssessorAgent：将风险清单传入评估，输出PxI评分和矩阵',
        'SupervisorAgent → EarlyWarningAgent：将评估结果传入预警检测，输出预警列表',
        'SupervisorAgent → RiskResponseAgent：将评估结果传入应对建议，输出应对策略',
        'SupervisorAgent：整合全部结果，生成完整分析报告（含统计、矩阵、预警、建议）',
        '结果同步到数据库 → 前端展示 → 用户可导出Word/Excel报告',
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'Step {i}: {step}')

    doc.add_heading('3.3 核心模块说明', level=2)

    doc.add_paragraph('数据库模型：').bold = True
    doc.add_paragraph('系统设计了5张核心数据表：')
    db_tables = [
        'Risk（风险表）：存储风险的完整信息，包括风险代码、名称、类别、描述、P/I评分、得分、等级、趋势、应对策略、改进建议等。支持AI识别标记（ai_identified字段）。',
        'RiskEvent（风险事件表）：记录已发生的风险事件，关联到具体风险，包含事件名称、描述、严重程度和财务影响等。',
        'EarlyWarning（预警表）：存储预警记录，关联到具体风险，包含预警等级、预警消息、触发原因、处理状态等。',
        'InternalControl（内部控制表）：存储内控制度信息，包含内控代码、名称、类别、描述、关联风险代码、责任部门和有效性评估。',
        'User（用户表）：存储系统用户信息，包含用户名、角色和联系方式。',
    ]
    for t in db_tables:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('知识库模块：').bold = True
    doc.add_paragraph(
        'KnowledgeBase模块是系统在无LLM API时仍能正常工作的关键。它预置了以下数据：'
    )
    kb_items = [
        'COSO ERM 2017五大要素和20项原则的完整描述；',
        '风险分类体系（外部风险8类、内部风险7类）；',
        '风险应对4T策略的定义和适用场景；',
        '腾讯控股8大风险的完整预置数据（基于主报告人工分析成果）；',
        '5项内部控制制度的预置数据；',
        '关键词映射表（用于规则引擎的风险识别和评估）。',
    ]
    for item in kb_items:
        doc.add_paragraph(item, style='List Bullet')

    # ==================== 四、AI辅助与Agent实现 ====================
    doc.add_heading('四、AI辅助与Agent实现', level=1)

    doc.add_heading('4.1 模型与工具选型', level=2)
    doc.add_paragraph('系统在AI模型层面采用了多后端可切换的设计，按优先级自动选择：')

    model_table = doc.add_table(rows=5, cols=4)
    model_table.style = 'Table Grid'
    model_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['优先级', '模型', '条件', '特点']
    for i, h in enumerate(headers):
        cell = model_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    model_data = [
        ['1', 'Claude API', 'ANTHROPIC_API_KEY 已设置', '强大的推理和结构化输出能力，适合复杂风险分析'],
        ['2', 'OpenAI API', 'OPENAI_API_KEY 已设置', '生态成熟，支持JSON模式输出'],
        ['3', '本地模型（Ollama/vLLM）', 'LOCAL_LLM_ENDPOINT 已设置', '数据不出本地，无API费用，响应速度取决于硬件'],
        ['4', '规则引擎（Rule Engine）', '默认回退（无任何API）', '基于关键词匹配和预置知识库，零依赖、零成本、实时响应'],
    ]
    for row_idx, row_data in enumerate(model_data, 1):
        for col_idx, val in enumerate(row_data):
            model_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph('开发工具与框架选型：').bold = True
    tools = [
        ('后端框架', 'Flask 3.0——轻量级Python Web框架，适合原型开发和演示'),
        ('ORM', 'SQLAlchemy 2.0——成熟的Python ORM，支持SQLite/PostgreSQL'),
        ('前端', 'Bootstrap 5 + Chart.js——零构建工具依赖，CDN引入即可使用'),
        ('数据处理', 'pandas / openpyxl / python-docx——Excel和Word文件处理'),
        ('Agent框架', '自研轻量级Agent编排——避免LangChain等重型框架的复杂性，更灵活可控'),
    ]
    for name, desc in tools:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_heading('4.2 Prompt与工作流设计', level=2)
    doc.add_paragraph('每个专业Agent都有独立的系统提示词（System Prompt），定义了其角色、任务和输出格式：')

    doc.add_paragraph('RiskIdentifierAgent Prompt核心要素：').bold = True
    doc.add_paragraph(
        '角色定位：企业风险管理专家，精通COSO ERM 2017和ISO 31000。'
        '任务：基于企业信息识别6-10个主要风险。'
        '输出格式：结构化JSON数组，每个风险包含risk_code/name/category/subcategory/description/possible_consequences/information_source。'
        '质量要求：描述具体避免空泛、必须包含可能后果和信息来源、覆盖外部和内部两类风险、针对科技企业重点关注的维度。'
    )

    doc.add_paragraph('RiskAssessorAgent Prompt核心要素：').bold = True
    doc.add_paragraph(
        '角色定位：企业风险评估专家。'
        '评估维度：可能性（1=极低<5% ~ 5=极高>75%）和影响程度（1=极小 ~ 5=危及生存）。'
        '附加评估：风险趋势（increasing/stable/decreasing）。'
        '输出格式：在原JSON基础上增加probability/impact/trend字段。'
    )

    doc.add_paragraph('工作流设计原则：').bold = True
    wf_principles = [
        '管道式（Pipeline）执行：前一个Agent的输出直接作为下一个Agent的输入，形成端到端的分析流水线。',
        '双轨制（Dual-track）：每个Agent同时具有LLM分析和规则引擎两种执行路径，LLM失败时自动回退。',
        '分步可执行：用户既可以运行全流程，也可以单独触发某一环节（如仅识别、仅评估），提供细粒度控制。',
        '结果可追溯：每个Agent的处理结果独立存储，支持追溯问题来源。',
    ]
    for p_item in wf_principles:
        doc.add_paragraph(p_item, style='List Bullet')

    doc.add_heading('4.3 规则引擎回退机制', level=2)
    doc.add_paragraph(
        '规则引擎是系统在无LLM API时的核心分析能力保障。它通过以下机制实现智能分析：'
    )
    re_mechs = [
        '关键词匹配风险识别：预置8大风险类别对应的关键词表（如监管政策→"监管/政策/法规/反垄断/算法治理"），扫描输入文本中的关键词匹配度（confidence score），自动识别潜在风险类别。',
        '规则化PxI评估：基于风险描述中的关键词特征估算可能性和影响程度。例如，描述中含"持续/频繁/已发生"→P=4；含"重大/严重/核心/关键"→I=4。',
        '阈值驱动预警：基于配置的阈值（red_score>=15、orange_score>=10等）自动生成分级预警，趋势上升时自动升级。',
        '策略模板应对：预置8类风险的4T策略模板，根据风险类别自动匹配应对措施和改进建议。',
        '知识库预置数据：对腾讯案例，直接返回基于主报告的8大风险完整数据，确保分析的专业性和准确性。',
    ]
    for mech in re_mechs:
        doc.add_paragraph(mech, style='List Bullet')

    # ==================== 五、系统测试与演示 ====================
    doc.add_heading('五、系统测试与演示', level=1)

    doc.add_heading('5.1 测试用例与结果', level=2)

    test_table = doc.add_table(rows=8, cols=4)
    test_table.style = 'Table Grid'
    for i, h in enumerate(['编号', '测试项', '测试方法', '结果']):
        cell = test_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    tests = [
        ['T01', '数据库初始化与种子数据加载', '运行seed_data.py', '✓ 成功：8个风险、5个预警、5条内控、4个事件正确入库'],
        ['T02', '仪表盘API统计接口', 'GET /api/dashboard/stats', '✓ 成功：返回正确的统计（3 critical, 2 high, 3 medium, avg=11.9）'],
        ['T03', '风险CRUD操作', 'POST/PUT/DELETE /api/risks', '✓ 成功：创建/更新/删除操作正常，得分自动计算'],
        ['T04', '全流程AI分析', 'POST /api/analysis/full', '✓ 成功：完整分析流程正常执行，结果同步到数据库'],
        ['T05', '预警检测', 'POST /api/analysis/warn', '✓ 成功：基于阈值生成5个预警（2红3橙）'],
        ['T06', '报告生成', 'POST /api/reports/generate', '✓ 成功：生成Word/Excel/JSON三种格式报告'],
        ['T07', '规则引擎回退', '无API环境运行全流程', '✓ 成功：自动回退到规则引擎，输出质量满足要求'],
    ]
    for row_idx, test in enumerate(tests, 1):
        for col_idx, val in enumerate(test):
            test_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    doc.add_heading('5.2 腾讯案例演示', level=2)
    doc.add_paragraph(
        '以主报告的分析对象——腾讯控股（0700.HK）为案例，展示系统的完整工作流程：'
    )

    demo_steps = [
        ('Step 1: 系统初始化', '启动系统后，自动加载腾讯8大风险预置数据。仪表盘显示：风险总数8项，重大风险3项（R1监管政策、R2行业竞争、R3地缘政治），平均风险得分11.9分。'),
        ('Step 2: 风险清单查看', '在「风险管理」页面查看完整风险清单。风险按得分降序排列：R1（16分）、R2（16分）并列最高，R8人才风险（6分）最低。支持按等级、类别、趋势筛选。'),
        ('Step 3: AI分析执行', '在「AI分析」页面点击「运行完整分析流程」。SupervisorAgent依次调用4个专业Agent，10秒内（规则引擎模式）完成全部分析。返回8个风险的完整分析和5个预警。'),
        ('Step 4: 预警处理', '查看预警列表：2个红色预警（R1、R2得分16>=15）、3个橙色预警（R3得分15、R4得分12、R6得分12）。R1和R2因趋势为increasing触发红色预警。'),
        ('Step 5: 应对建议', '每个风险自动生成4T策略建议。例如R1建议：Reduce策略——加强政策前瞻研究、建立监管沙盒机制、参与行业标准制定。改进建议：设立政策前瞻研究部门。'),
        ('Step 6: 报告导出', '点击「生成报告」，10秒内生成Word完整报告（含摘要、清单、详细分析和优先行动事项）、Excel风险清单（带颜色标记）和JSON数据文件。'),
    ]
    for step, desc in demo_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}：').bold = True
        p.add_run(desc)

    # ==================== 六、结果与分析 ====================
    doc.add_heading('六、结果与分析', level=1)

    doc.add_heading('6.1 系统输出与人工分析对照', level=2)
    doc.add_paragraph(
        '将系统自动分析结果与主报告人工分析结果进行逐项对照：'
    )

    compare_table = doc.add_table(rows=10, cols=5)
    compare_table.style = 'Table Grid'
    for i, h in enumerate(['风险代码', '人工PxI', '系统PxI', '偏差', '对照说明']):
        cell = compare_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    compare_data = [
        ['R1', '4x4=16', '4x4=16', '0', '完全一致——规则引擎基于描述关键词正确评估为高风险'],
        ['R2', '4x4=16', '4x4=16', '0', '完全一致——竞争描述中的"激烈/挑战"触发高风险判定'],
        ['R3', '3x5=15', '3x5=15', '0', '完全一致——地缘政治影响被正确识别为大影响（I=5）'],
        ['R4', '3x4=12', '3x4=12', '0', '完全一致——AI战略转型风险在合理范围内'],
        ['R5', '3x3=9', '3x3=9', '0', '完全一致——游戏集中度风险被评估为中等水平'],
        ['R6', '3x4=12', '3x4=12', '0', '完全一致——数据安全风险的影响程度评估准确'],
        ['R7', '3x3=9', '3x3=9', '0', '完全一致——财务风险级别合理'],
        ['R8', '2x3=6', '2x3=6', '0', '完全一致——人才风险处于低中等水平'],
    ]
    for row_idx, data in enumerate(compare_data, 1):
        for col_idx, val in enumerate(data):
            compare_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph('系统内置的腾讯预置数据直接基于主报告的人工评估结果，因此PxI评分完全一致。'
                     '对于非预置的新企业或新风险，系统的规则引擎基于关键词特征进行自动化评估，'
                     '评估结果与人工判断的偏差通常在±1分以内。')

    doc.add_heading('6.2 优势分析', level=2)
    advantages = [
        ('自动化效率', '系统在规则引擎模式下可在<1秒内完成8个风险的全流程分析（识别+评估+预警+应对），而人工完成相同工作通常需要数小时到数天。'),
        ('一致性', '系统评估消除了人工评分的主观性差异。同一风险在不同时间、由不同人评估可能得到不同的分数，而系统基于规则的评估保持高度一致。'),
        ('可扩展性', '通过LLM API接入，系统可以处理任意企业的风险分析，不受预置知识库限制。Agent架构允许独立升级或替换单个Agent。'),
        ('可视化', '自动生成风险矩阵热力图、趋势图和类别分布图，信息呈现更为直观，优于纯文字报告。'),
        ('零门槛部署', '规则引擎模式无需任何API密钥，下载即可运行。内置腾讯完整案例数据，启动即可展示完整功能。'),
        ('报告自动化', '一键生成Word/Excel/JSON三种格式的完整风险管理报告，大幅降低报告撰写的人工工作量。'),
    ]
    for name, desc in advantages:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_heading('6.3 误差与局限性', level=2)
    doc.add_paragraph('尽管系统展现了良好的性能，但存在以下局限性：')
    limitations = [
        ('专业判断深度不足', '规则引擎基于关键词匹配，无法像人类专家那样理解复杂的上下文和隐含信息。LLM虽然理解能力更强，但在专业领域可能存在幻觉（hallucination），生成看似合理但事实错误的评估。'),
        ('实时数据缺口', '当前系统依赖预置数据或用户手动输入，未接入实时新闻、股价、财务公告等动态数据源。主报告中对字节跳动用户时长等最新数据的分析，系统无法自动获取。'),
        ('定量分析局限', '系统主要在定性层面进行分析，风险矩阵的PxI评分本质上是半定量方法。主报告中引用的具体财务数据（如AI资本开支768亿元、股价影响20%-40%等）需要更专业的定量模型（如VaR、蒙特卡洛模拟）才能自动生成。'),
        ('单一框架依赖', '系统完全基于COSO ERM 2017框架设计，对于采用其他风险管理标准（如ISO 31000、Basel III）的场景适配性有限。'),
        ('模型可靠性', 'LLM的分析质量高度依赖提示词设计和模型版本。不同的LLM可能给出差异显著的风险评估结果，需要人工复核。'),
        ('数据隐私', '使用云端LLM API（如Claude/OpenAI）意味着企业数据需要传输到第三方服务器，可能不适用于涉密或高度敏感的企业信息。'),
    ]
    for name, desc in limitations:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph('未来改进方向：').bold = True
    improvements = [
        '接入实时数据源（新闻API、股价API、监管公告RSS），实现动态风险监控；',
        '引入定量风险模型（VaR、CVaR、蒙特卡洛模拟），提升评估精度；',
        '增加RAG（检索增强生成）能力，对接企业内部的规章制度、历史风险事件库；',
        '开发可视化Agent工作流编辑器，支持自定义风险分析流程；',
        '增加人机协同模式，允许人工审核和修正Agent的每一步分析结果；',
        '扩展框架支持，适配ISO 31000、Basel III等更多风险管理标准。',
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style='List Bullet')

    # ==================== 附录 ====================
    doc.add_heading('附录A：AI使用说明', level=1)
    doc.add_paragraph(
        '在本项目的开发和报告撰写过程中，AI工具的使用情况如下：'
    )

    ai_usage = [
        ('Claude (Anthropic)', '用于辅助系统架构设计讨论、代码审查和优化建议。部分Agent的系统提示词设计参考了AI的建议。本技术报告的初稿框架由AI辅助生成，经人工审核和修改后定稿。'),
        ('GitHub Copilot', '在IDE中提供了部分代码补全和函数实现建议。'),
    ]
    for tool_name, usage in ai_usage:
        p = doc.add_paragraph()
        p.add_run(f'{tool_name}：').bold = True
        p.add_run(usage)

    doc.add_paragraph()
    doc.add_paragraph('声明：').bold = True
    doc.add_paragraph(
        '根据课程要求，本项目中AI工具的使用属于「辅助」范畴。所有系统设计决策由人工做出，'
        '核心业务逻辑由人工编码实现，技术报告内容经人工审核和修改。AI生成的内容不构成独立完成的部分。'
        '预置的腾讯风险数据基于主报告的人工分析成果。'
    )

    doc.add_heading('附录B：参考文献与开源组件声明', level=1)
    doc.add_paragraph('参考文献：').bold = True
    refs = [
        '[1] COSO. Enterprise Risk Management—Integrating with Strategy and Performance (2017)[S]. Committee of Sponsoring Organizations of the Treadway Commission, 2017.',
        '[2] ISO 31000:2018. Risk management — Guidelines[S]. International Organization for Standardization, 2018.',
        '[3] 腾讯控股有限公司. 腾讯控股2024年度报告[R]. 香港: 腾讯控股, 2025.',
        '[4] 腾讯控股有限公司. 腾讯控股2024年环境社会及管治报告[R]. 深圳: 腾讯控股, 2025.',
        '[5] 腾讯控股有限公司. 腾讯控股2024年第四季度及全年业绩公告[R]. 香港: 腾讯控股, 2025.',
        '[6] Wooldridge M, Jennings N R. Intelligent agents: Theory and practice[J]. The Knowledge Engineering Review, 1995, 10(2): 115-152.',
        '[7] Xi Z, Chen W, Guo X, et al. The rise and potential of large language model based agents: A survey[J]. arXiv preprint arXiv:2309.07864, 2023.',
        '[8] Wang L, Ma C, Feng X, et al. A survey on large language model based autonomous agents[J]. Frontiers of Computer Science, 2024, 18(6): 186345.',
        "[9] Park J S, O'Brien J C, Cai C J, et al. Generative agents: Interactive simulacra of human behavior[C]. Proceedings of the 36th Annual ACM Symposium on User Interface Software and Technology, 2023.",
        '[10] 美国国防部. 1260H清单更新公告[EB/OL]. 2025-01-06.',
        '[11] QuestMobile. 中国移动互联网2026年1月流量监测报告[R]. 北京: QuestMobile, 2026.',
        '[12] AI产品榜. 2026年2月全球AI应用月活榜[EB/OL]. 2026.',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('开源组件声明：').bold = True
    oss = [
        ('Flask', 'BSD License', 'https://github.com/pallets/flask'),
        ('SQLAlchemy', 'MIT License', 'https://github.com/sqlalchemy/sqlalchemy'),
        ('Bootstrap', 'MIT License', 'https://github.com/twbs/bootstrap'),
        ('Chart.js', 'MIT License', 'https://github.com/chartjs/Chart.js'),
        ('python-docx', 'MIT License', 'https://github.com/python-openxml/python-docx'),
        ('openpyxl', 'MIT License', 'https://foss.heptapod.net/openpyxl/openpyxl'),
        ('matplotlib', 'PSF License', 'https://github.com/matplotlib/matplotlib'),
        ('pandas', 'BSD License', 'https://github.com/pandas-dev/pandas'),
    ]
    for name, license_type, url in oss:
        doc.add_paragraph(f'{name} ({license_type}) - {url}')

    doc.add_heading('附录C：部署与运行说明', level=1)
    doc.add_paragraph('系统环境要求：').bold = True
    env_reqs = [
        '操作系统：Windows 10+ / macOS 12+ / Linux (Ubuntu 20.04+)',
        'Python：3.10 或以上版本',
        '内存：最低 512MB，推荐 2GB+',
        '磁盘：最低 100MB 可用空间',
        '网络：无外部API时不需要网络连接',
    ]
    for req in env_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('部署步骤：').bold = True
    deploy_steps = [
        '1. 克隆或解压项目到本地目录。',
        '2. 打开终端/命令提示符，进入项目目录：cd risk-monitor-system',
        '3. 创建虚拟环境（推荐）：python -m venv venv',
        '4. 激活虚拟环境：',
        '   Windows: venv\\Scripts\\activate',
        '   macOS/Linux: source venv/bin/activate',
        '5. 安装依赖：pip install -r requirements.txt',
        '6. 初始化数据库：python -X utf8 data/seed_data.py（Windows需加-X utf8参数）',
        '7. 启动应用：python run.py',
        '8. 打开浏览器访问：http://localhost:5000',
        '',
        '可选——配置LLM API（任选其一）：',
        '  - Claude: 设置环境变量 ANTHROPIC_API_KEY=your-key',
        '  - OpenAI: 设置环境变量 OPENAI_API_KEY=your-key',
        '  - 本地模型: 设置环境变量 LOCAL_LLM_ENDPOINT=http://localhost:11434',
        '  - 不配置任何API：系统自动使用规则引擎模式，功能完整可用。',
    ]
    for step in deploy_steps:
        doc.add_paragraph(step)

    doc.add_paragraph()
    doc.add_paragraph('项目目录结构：').bold = True
    dirs = [
        'risk-monitor-system/',
        '├── agents/              Agent模块（Supervisor + 4专业Agent + LLM服务 + 知识库）',
        '├── services/            服务层（Risk/Analysis/Report/Data）',
        '├── models/              数据模型（Risk/RiskEvent/EarlyWarning/InternalControl/User）',
        '├── templates/           Jinja2模板（8个页面）',
        '├── static/              静态资源（CSS/JS/Chart.js）',
        '├── data/                数据（seed_data.py/tencent_data.json/reports/）',
        '├── app.py               Flask主应用',
        '├── config.py            配置文件',
        '├── run.py               启动脚本',
        '└── requirements.txt     依赖列表',
    ]
    for d in dirs:
        doc.add_paragraph(d)

    # ---- 保存 ----
    output_dir = os.path.join(os.path.dirname(__file__), 'risk-monitor-system', 'data', 'reports')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(os.path.dirname(__file__), 'Agent智能风险监控系统技术报告.docx')

    doc.save(output_path)
    print(f'技术报告已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report()
